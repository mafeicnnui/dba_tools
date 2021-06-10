#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time : 2020/11/19 13:59
# @Author : 马飞
# @File : py2word.py
# @Software: PyCharm

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('商管系统运行周报', 0)  #插入标题

p = document.add_paragraph('A plain paragraph having some ')   #插入段落
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('tu.jpg', width=Inches(1.25)) #插入图片

table = document.add_table(rows=1, cols=3) #插入表格
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# for item in hdr_cells:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(item.qty)
#     row_cells[1].text = str(item.id)
#     row_cells[2].text = item.desc

document.add_page_break()

document.save('demo.docx')  #保存文档